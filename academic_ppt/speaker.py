"""Create editable DOCX speaker scripts from confirmed page plans."""

from __future__ import annotations

from pathlib import Path

from .planning import PagePlan


class SpeakerScriptWriter:
    """Write a presenter-facing script without exposing production audit details."""

    def write(self, plan: PagePlan, output_path: Path | str) -> Path:
        try:
            from docx import Document
        except ImportError as exc:
            raise RuntimeError("python-docx is required to write speaker scripts") from exc
        output = Path(output_path)
        if output.suffix.lower() != ".docx":
            raise ValueError("speaker script must use a .docx path")
        output.parent.mkdir(parents=True, exist_ok=True)
        document = Document()
        document.add_heading(f"{plan.scene} 演讲稿", level=0)
        for page in plan.pages:
            document.add_heading(f"{page.page_id} {page.title}", level=1)
            document.add_paragraph(f"本页回答：{page.question_answered}")
            document.add_paragraph(f"讲解：{page.interpretation}")
            document.add_paragraph(f"转场：{page.next_link}")
            document.add_paragraph(f"建议时长：{page.time_seconds} 秒")
        document.save(output)
        return output.resolve()
